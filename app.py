import os
import re
import time
import shutil
import zipfile
import unicodedata
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from copy import deepcopy

import pandas as pd
from flask import Flask, request, render_template_string, send_file, jsonify

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024

# =====================
# PATHS
# =====================
BASE_DIR = Path(__file__).resolve().parent
UPLOAD_FOLDER  = BASE_DIR / "uploads"
OUTPUT_FOLDER  = BASE_DIR / "output_grouped"
ARCHIVE_FOLDER = BASE_DIR / "archive_data"
WORD_TEMPLATE  = BASE_DIR / "template_chung_tu_clean.docx"
EXCEL_DATA     = BASE_DIR / "TEMPLATE_RUN_MEGER_-_Copy.xlsx"
TICKET_LIST    = BASE_DIR / "freshdesk_ticket_list.xlsx"
FRESHDESK_AUTO = BASE_DIR / "freshdesk_reply_list_auto.xlsx"
LOG_FILE       = BASE_DIR / "freshdesk_reply_log.xlsx"

for d in [UPLOAD_FOLDER, OUTPUT_FOLDER, ARCHIVE_FOLDER]:
    d.mkdir(exist_ok=True)

# =====================
# MOCK FRESHDESK DATABASE
# Tao tu file freshdesk_ticket_list.xlsx + reply list
# =====================
def build_mock_freshdesk_db():
    db = {}
    # Load ticket IDs tu file
    if TICKET_LIST.exists():
        df = pd.read_excel(TICKET_LIST)
        for tid in df["ticket_id"].astype(str).tolist():
            db[tid.strip()] = {
                "id": tid.strip(),
                "subject": f"[Mock] Ticket #{tid.strip()} - Khiếu nại giao dịch",
                "status": "open",
                "cc_emails": [f"bank_mock_{tid[-4:]}@mockbank.com"],
                "replies": []
            }
    # Load them tu reply list
    if FRESHDESK_AUTO.exists():
        df2 = pd.read_excel(FRESHDESK_AUTO)
        for _, row in df2.iterrows():
            tid = str(row.get("ticket_id", "")).strip()
            if tid and tid not in db:
                db[tid] = {
                    "id": tid,
                    "subject": f"[Mock] Ticket #{tid} - {row.get('file_name','')}",
                    "status": "open",
                    "cc_emails": [f"bank_mock@mockbank.com"],
                    "replies": []
                }
    return db

MOCK_DB = build_mock_freshdesk_db()

# =====================
# COLUMN CONFIG
# =====================
GROUP_BY_COLUMN = "User ID"
FILE_NAME_PATTERN = "CHUNG TU {Bank}_ZALOPAY_{ten_file}.docx"

COLUMN_MAP = {
    "Bank": "Bank",
    "tên file": "ten_file",
    "User ID": "user_id",
    "BC Trace No": "bc_trace_no",
    "TPE Trans ID": "tpe_trans_id",
    "App ID": "app_id",
    "Card Number": "card_number",
    "Trans Time": "trans_time",
    "Amount": "amount",
    "Sub Trans Type": "sub_trans_type",
    "Status (BC)": "bc_status",
    "Phone": "phone",
    "Connector Channel": "conn_channel",
    "App User": "app_user",
    "Note": "note",
    "so_du": "so_du",
    "ngay_tao_vi": "ngay_tao_vi",
    "ngay_lien_ket": "ngay_lien_ket",
    "lien_ket_lan_dau": "lien_ket_lan_dau",
}

TRANSACTION_PLACEHOLDERS = {
    "bc_trace_no", "tpe_trans_id", "app_id", "card_number", "trans_time",
    "amount", "sub_trans_type", "bc_status", "phone", "conn_channel", "app_user", "note"
}

ADD_CC_EMAILS = ["banksupport@vng.com.vn", "chargebackzp@vng.com.vn"]
REMOVE_DOMAINS = {"vng.com.vn", "zalopay.vn"}

# =====================
# HELPERS
# =====================
def clean_value(value):
    if pd.isna(value):
        return ""
    if hasattr(value, "strftime"):
        if getattr(value, "hour", 0) or getattr(value, "minute", 0) or getattr(value, "second", 0):
            return value.strftime("%d/%m/%Y %H:%M:%S")
        return value.strftime("%d/%m/%Y")
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()

def safe_filename(name):
    name = re.sub(r'[\\/:*?"<>|]+', "_", str(name))
    return re.sub(r"\s+", " ", name).strip() or "NO_NAME"

def is_six_digit(value):
    return bool(re.fullmatch(r"\d{6}", clean_value(value)))

def build_doc_title(ctx):
    bank = ctx.get("Bank", "")
    tf   = ctx.get("ten_file", "")
    card = re.sub(r"\s+", "", clean_value(ctx.get("card_number", "")))
    return f"CHUNG TU {bank}_ZALOPAY.{tf}_{card}" if is_six_digit(tf) else f"CHUNG TU {bank}_ZALOPAY.{tf}"

def build_output_filename(ctx):
    if is_six_digit(ctx.get("ten_file", "")):
        return build_doc_title(ctx) + ".docx"
    return FILE_NAME_PATTERN.format(**ctx)

def normalize_text(text):
    text = clean_value(text).lower()
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", text).strip()

def parse_money(value):
    if pd.isna(value):
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    text = re.sub(r"[^\d.-]", "", str(value).replace(",", ""))
    try:
        return float(text)
    except Exception:
        return 0.0

def norm_series(group_df, col):
    if col not in group_df.columns:
        return []
    return [normalize_text(v) for v in group_df[col].fillna("").astype(str)]

def get_service_line1(group_df):
    app_ids = [v for v in norm_series(group_df, "App ID") if v]
    if not app_ids:
        return "Giao dịch thanh toán đơn hàng qua ví Zalopay bằng thẻ/ tài khoản ngân hàng đã liên kết."
    google  = normalize_text("579 - Google Play")
    apple   = normalize_text("9999 - Apple Services")
    topup   = normalize_text("454 - Nạp tiền")
    if all(v == google for v in app_ids):
        return "Giao dịch thanh toán phí ứng dụng/ phí dịch vụ Google Play qua phương thức ví Zalopay bằng thẻ/ tài khoản ngân hàng đã liên kết."
    if all(v == apple for v in app_ids):
        return "Giao dịch thanh toán phí ứng dụng/ phí dịch vụ Apple qua phương thức ví Zalopay bằng thẻ/ tài khoản ngân hàng đã liên kết."
    if all(v == topup for v in app_ids):
        return "Giao dịch nạp tiền vào ví Zalopay bằng thẻ/ tài khoản ngân hàng đã liên kết."
    if any(v == topup for v in app_ids) and any(v != topup for v in app_ids):
        return "Giao dịch thanh toán đơn hàng và nạp tiền vào ví Zalopay bằng thẻ/ tài khoản ngân hàng đã liên kết."
    return "Giao dịch thanh toán đơn hàng qua ví Zalopay bằng thẻ/ tài khoản ngân hàng đã liên kết."

def get_service_line2(group_df):
    statuses = [v for v in norm_series(group_df, "Status (BC)") if v]
    if not statuses:
        return ""
    refund  = sum("hoan tien" in v for v in statuses)
    success = sum("successful" in v for v in statuses)
    ok = "Giao dịch thanh toán thành công, đã cung cấp hàng hóa/ dịch vụ."
    if success > 0 and refund > 0:
        return f"{ok} {refund} giao dịch đã hoàn tiền."
    if refund == len(statuses):
        return "Giao dịch đã hoàn tiền."
    if success > 0:
        return ok
    return ""

def get_wallet_note(row):
    val = row.get("so_du", row.get("so_du", ""))
    if parse_money(val) > 100000:
        return (
            "Trong ví vẫn có một phần số dư có thể hỗ trợ thu hồi. "
            "Nếu cần thu hồi anh/chị vui lòng cung cấp công văn thu hồi hoặc thư khiếu nại của KH "
            "(bản chính hoặc bản sao có dấu) gửi qua bưu điện cho Zalopay để Zalopay hỗ trợ tiếp.\n"
            "Địa chỉ nhận thư như chữ ký email này - vui lòng ghi rõ tên người nhận. "
            "Anh chị lưu ý cung cấp số tài khoản còn hoạt động để Zalopay có thể xử lý khi hoàn trả."
        )
    return ""
def build_context(row):
    ctx = {}
    for excel_col, ph in COLUMN_MAP.items():
        val = clean_value(row.get(excel_col, ""))
        if ph == "amount" and val:
            try:
                val = f"{float(str(val).replace(',','')):.0f}"
            except Exception:
                pass
        if ph == "sub_trans_type":
            val = re.sub(r"\s*\(([^)]+)\)\s*$", r"\n(\1)", val)
        if ph == "app_id":
            val = re.sub(r"\s*-\s*", " - ", val)
        ctx[ph] = val
    return ctx

def remove_mail_merge(docx_path):
    import xml.etree.ElementTree as ET
    docx_path = Path(docx_path)
    tmp = docx_path.with_suffix(".tmp")
    ET.register_namespace("", "http://schemas.openxmlformats.org/package/2006/relationships")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                text = data.decode("utf-8", errors="ignore")
                text = re.sub(r"<w:mailMerge[\s\S]*?</w:mailMerge>", "", text)
                data = text.encode("utf-8")
            elif item.filename == "word/_rels/settings.xml.rels":
                try:
                    root = ET.fromstring(data)
                    for rel in list(root):
                        if "mailMergeSource" in rel.attrib.get("Type", ""):
                            root.remove(rel)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                except Exception:
                    pass
            zout.writestr(item, data)
    tmp.replace(docx_path)

def replace_para(para, ctx):
    full = para.text
    new  = full
    for k, v in ctx.items():
        new = new.replace("{{" + k + "}}", str(v))
    if new != full and para.runs:
        for r in para.runs:
            r.text = ""
        para.runs[0].text = new

def replace_doc(doc, ctx):
    for p in doc.paragraphs:
        replace_para(p, ctx)
    for tbl in doc.tables:
        txt = "\n".join(c.text for row in tbl.rows for c in row.cells)
        if any("{{" + ph + "}}" in txt for ph in TRANSACTION_PLACEHOLDERS):
            continue
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_para(p, ctx)
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            replace_para(p, ctx)
        for p in sec.footer.paragraphs:
            replace_para(p, ctx)

def render_rows(doc, group_df):
    table = None
    for tbl in doc.tables:
        txt = "\n".join(c.text for row in tbl.rows for c in row.cells)
        if any("{{" + ph + "}}" in txt for ph in TRANSACTION_PLACEHOLDERS):
            table = tbl
            break
    if table is None:
        return
    tmpl_row = None
    for row in table.rows:
        if any("{{" + ph + "}}" in "\n".join(c.text for c in row.cells) for ph in TRANSACTION_PLACEHOLDERS):
            tmpl_row = row
            break
    if tmpl_row is None:
        return
    tmpl_tr = deepcopy(tmpl_row._tr)
    tmpl_row._tr.getparent().remove(tmpl_row._tr)
    for _, row in group_df.iterrows():
        new_tr = deepcopy(tmpl_tr)
        table._tbl.append(new_tr)
        ctx = build_context(row)
        for cell in table.rows[-1].cells:
            for p in cell.paragraphs:
                replace_para(p, ctx)

# =====================
# MOCK FRESHDESK API
# =====================
@app.route("/mock-freshdesk/api/v2/tickets/<ticket_id>", methods=["GET"])
def mock_get_ticket(ticket_id):
    ticket = MOCK_DB.get(str(ticket_id))
    if not ticket:
        return jsonify({"error": "Ticket not found"}), 404
    return jsonify({
        "id": ticket["id"],
        "subject": ticket["subject"],
        "status": 2,
        "cc_emails": ticket["cc_emails"],
        "description": f"Mock ticket #{ticket_id}"
    })

@app.route("/mock-freshdesk/api/v2/tickets/<ticket_id>/reply", methods=["POST"])
def mock_reply_ticket(ticket_id):
    ticket = MOCK_DB.get(str(ticket_id))
    if not ticket:
        return jsonify({"error": "Ticket not found"}), 404
    body = request.form.get("body", "")
    cc_emails = request.form.getlist("cc_emails[]")
    has_attachment = "attachments[]" in request.files
    ticket["replies"].append({
        "body": body[:100],
        "cc_emails": cc_emails,
        "has_attachment": has_attachment,
        "timestamp": datetime.now().isoformat()
    })
    return jsonify({
        "id": f"reply_{ticket_id}_{len(ticket['replies'])}",
        "ticket_id": ticket_id,
        "body": body[:100],
        "cc_emails": cc_emails,
        "attachments": [{"name": "file.docx"}] if has_attachment else []
    }), 201

@app.route("/mock-freshdesk/tickets", methods=["GET"])
def mock_ticket_list():
    tickets = []
    for tid, t in MOCK_DB.items():
        tickets.append({
            "id": t["id"],
            "subject": t["subject"],
            "status": "open",
            "cc_emails": t["cc_emails"],
            "replies_count": len(t["replies"])
        })
    return jsonify({"tickets": tickets, "total": len(tickets)})

# =====================
# APP STATE
# =====================
state = {"files": [], "freshdesk_data": [], "send_results": []}

# =====================
# HTML
# =====================
HTML = """<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>ZOP_HR_INBOX</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',Arial,sans-serif;background:#f0f2f5;color:#1a1a2e;min-height:100vh}
/* ── Topbar ── */
.topbar{background:#003087;color:white;padding:12px 32px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:100;box-shadow:0 2px 8px rgba(0,0,40,.18)}
.topbar h1{font-size:17px;font-weight:600}
.topbar .right{display:flex;align-items:center;gap:10px;font-size:13px}
/* ── Stepper ── */
.stepper{background:white;border-bottom:1px solid #e5e7eb;padding:0 32px;position:sticky;top:49px;z-index:99;box-shadow:0 1px 4px rgba(0,0,0,.05)}
.stepper-inner{max-width:960px;margin:0 auto;display:flex;align-items:stretch}
.step-item{display:flex;align-items:center;gap:10px;padding:14px 20px 14px 0;flex:1;cursor:default;border-bottom:3px solid transparent;transition:.2s;color:#94a3b8;font-size:13px;font-weight:500}
.step-item:first-child{padding-left:0}
.step-item.done{color:#15803d;cursor:pointer}
.step-item.done:hover .s-num{background:#d1fae5;color:#15803d}
.step-item.active{color:#003087;border-bottom-color:#003087}
.s-num{width:26px;height:26px;border-radius:50%;background:#e5e7eb;color:#94a3b8;display:inline-flex;align-items:center;justify-content:center;font-size:12px;font-weight:700;flex-shrink:0;transition:.2s}
.step-item.active .s-num{background:#003087;color:white}
.step-item.done .s-num{background:#dcfce7;color:#15803d}
.step-sep{color:#cbd5e1;font-size:18px;padding:0 4px;align-self:center}
/* ── Page layout ── */
.wrap{max-width:960px;margin:0 auto;padding:28px 16px}
.page{display:none}.page.active{display:block}
/* ── Card ── */
.card{background:white;border-radius:14px;border:1px solid #e5e7eb;padding:24px 28px;margin-bottom:18px;box-shadow:0 1px 4px rgba(0,0,0,.06)}
.card-title{font-size:15px;font-weight:600;margin-bottom:18px;color:#003087;display:flex;align-items:center;gap:10px}
/* ── Shared ── */
.badge{display:inline-block;padding:3px 10px;border-radius:20px;font-size:12px;font-weight:600}
.badge-mock{background:#fbbf24;color:#78350f}
.badge-ok{background:#86efac;color:#14532d}
.badge-skip{background:#e5e7eb;color:#6b7280}
.badge-err{background:#fca5a5;color:#7f1d1d}
.drop-zone{border:2px dashed #cbd5e1;border-radius:10px;padding:28px;text-align:center;cursor:pointer;background:#f8fafc;transition:.2s}
.drop-zone:hover{border-color:#003087;background:#eff6ff}
.drop-zone input{display:none}
.drop-icon{font-size:36px;margin-bottom:8px}
.drop-text{color:#64748b;font-size:14px}
.drop-text b{color:#003087}
.btn{display:inline-flex;align-items:center;gap:7px;padding:9px 18px;border-radius:8px;font-size:13px;font-weight:500;cursor:pointer;border:none;transition:.15s}
.btn-blue{background:#003087;color:white}.btn-blue:hover{background:#002060}
.btn-blue:disabled,.btn-green:disabled,.btn-gray:disabled{background:#9ca3af!important;cursor:not-allowed!important;color:white!important}
.btn-green{background:#15803d;color:white}.btn-green:hover{background:#166534}
.btn-gray{background:#f1f5f9;color:#334155;border:1px solid #cbd5e1}.btn-gray:hover{background:#e2e8f0}
.btn-ghost{background:none;border:none;color:#64748b;font-size:13px;cursor:pointer;padding:9px 4px;display:inline-flex;align-items:center;gap:5px}
.btn-ghost:hover{color:#1e40af}
.progress{margin-top:12px;display:none}
.bar{height:7px;background:#e5e7eb;border-radius:4px;overflow:hidden}
.fill{height:100%;background:#003087;width:0%;transition:width .3s;border-radius:4px}
.ptext{font-size:12px;color:#64748b;margin-top:5px}
table{width:100%;border-collapse:collapse;font-size:13px}
th{background:#f8fafc;padding:9px 11px;text-align:left;font-weight:600;border-bottom:2px solid #e5e7eb;color:#374151}
td{padding:9px 11px;border-bottom:1px solid #f1f5f9;color:#1f2937}
tr:hover td{background:#f8fafc}
.alert{padding:11px 15px;border-radius:8px;font-size:13px;margin-bottom:14px}
.alert-info{background:#eff6ff;border:1px solid #bfdbfe;color:#1e40af}
.alert-success{background:#f0fdf4;border:1px solid #86efac;color:#166534}
.alert-warn{background:#fffbeb;border:1px solid #fde68a;color:#92400e}
.actions{display:flex;gap:8px;flex-wrap:wrap;margin-top:16px;align-items:center}
.actions-between{justify-content:space-between}
.preview-email{background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;padding:14px;font-size:13px;line-height:1.7;margin-bottom:10px;position:relative}
.preview-email .meta{color:#64748b;font-size:11px;border-bottom:1px solid #e2e8f0;padding-bottom:8px;margin-bottom:10px}
.preview-attach{display:inline-flex;align-items:center;gap:5px;background:#e0f2fe;color:#0369a1;padding:4px 10px;border-radius:6px;font-size:12px;margin-top:8px}
.ticket-row.unchecked{opacity:.4}
.ticket-checkbox{position:absolute;top:12px;right:12px;width:18px;height:18px;cursor:pointer;accent-color:#003087}
.sel-bar{display:flex;align-items:center;gap:10px;margin-bottom:12px;font-size:13px;flex-wrap:wrap}
.sel-count{font-weight:600;color:#003087}
.tab-bar{display:flex;gap:2px;margin-bottom:16px;border-bottom:2px solid #e5e7eb}
.tab{padding:8px 16px;cursor:pointer;font-size:13px;border-radius:6px 6px 0 0;color:#64748b;border:1px solid transparent;border-bottom:none}
.tab.active{background:white;border-color:#e5e7eb;color:#003087;font-weight:600;margin-bottom:-2px}
.tab-content{display:none}.tab-content.active{display:block}
.stat-row{display:grid;grid-template-columns:repeat(3,1fr);gap:12px;margin-bottom:16px}
.stat{background:#f8fafc;border-radius:8px;padding:12px;text-align:center}
.stat .num{font-size:26px;font-weight:700;color:#003087}
.stat .lbl{font-size:12px;color:#64748b;margin-top:2px}
</style>
</head>
<body>

<!-- Topbar -->
<div class="topbar">
  <h1>📬 ZOP_HR_INBOX</h1>
  <div class="right">
    <span class="badge badge-mock">🧪 MOCK MODE</span>
    <span style="font-size:12px">Mock Freshdesk: <b>localhost/mock-freshdesk</b></span>
  </div>
</div>

<!-- Stepper -->
<div class="stepper">
  <div class="stepper-inner">
    <div class="step-item active" id="si-1"><span class="s-num" id="sn-1">1</span> Upload Excel</div>
    <span class="step-sep">›</span>
    <div class="step-item" id="si-2"><span class="s-num" id="sn-2">2</span> Kết quả Word</div>
    <span class="step-sep">›</span>
    <div class="step-item" id="si-3"><span class="s-num" id="sn-3">3</span> Chọn &amp; Gửi</div>
    <span class="step-sep">›</span>
    <div class="step-item" id="si-4"><span class="s-num" id="sn-4">4</span> Hoàn thành</div>
  </div>
</div>

<!-- Pages -->
<div class="wrap">

  <!-- PAGE 1 -->
  <div class="page active" id="pg-1">
    <div class="card">
      <div class="card-title">📊 Upload file Excel giao dịch</div>
      <div class="alert alert-info" style="margin-bottom:16px">
        ℹ️ <b>Demo Mode:</b> File Word tạo thật từ data giả lập. Freshdesk được giả lập bằng Mock API — không gửi email thật.
      </div>
      <div class="drop-zone" onclick="document.getElementById('fi').click()">
        <input type="file" id="fi" accept=".xlsx,.xls" onchange="onFileSelect(this)">
        <div class="drop-icon">📊</div>
        <div class="drop-text"><b>Bấm để chọn file Excel</b><br>Hỗ trợ .xlsx · Tối đa 32MB</div>
      </div>
      <div id="fname" style="font-size:13px;color:#003087;margin-top:8px"></div>
      <div class="actions">
        <button class="btn btn-blue" id="btn1" onclick="processExcel()" disabled>⚙️ Tạo file Word chứng từ</button>
      </div>
      <div class="progress" id="p1">
        <div class="bar"><div class="fill" id="f1"></div></div>
        <div class="ptext" id="pt1">Đang xử lý...</div>
      </div>
    </div>
  </div>

  <!-- PAGE 2 -->
  <div class="page" id="pg-2">
    <div class="card">
      <div class="card-title">📄 Kết quả tạo file Word chứng từ</div>
      <div id="s2body"></div>
      <div class="actions actions-between" id="s2act"></div>

      <!-- Convert progress -->
      <div id="convert-progress" style="display:none;margin-top:14px">
        <div class="bar"><div class="fill" id="convert-fill" style="width:0%"></div></div>
        <div class="ptext" id="convert-text">⏳ Đang chuyển đổi sang PDF...</div>
      </div>

      <!-- Upload thay thế -->
      <div id="replace-zone" style="display:none;margin-top:22px;border-top:1px solid #e5e7eb;padding-top:18px">
        <div style="font-size:14px;font-weight:600;color:#92400e;margin-bottom:10px">📝 Điều chỉnh chứng từ <span style="font-weight:400;font-size:12px;color:#78716c">(tùy chọn)</span></div>
        <div class="alert alert-warn" style="margin-bottom:12px">
          Tải ZIP → Sửa file Word → Upload lại đây (giữ nguyên tên file) → Tiếp tục gửi.
        </div>
        <div class="drop-zone" id="replace-drop"
             ondragover="event.preventDefault();this.style.borderColor='#ea580c'"
             ondragleave="this.style.borderColor='#cbd5e1'"
             ondrop="handleReplaceDrop(event)"
             onclick="document.getElementById('replace-fi').click()"
             style="padding:18px">
          <input type="file" id="replace-fi" accept=".docx" multiple onchange="handleReplaceSelect(this)">
          <div class="drop-icon" style="font-size:28px">📤</div>
          <div class="drop-text"><b>Upload file .docx thay thế</b><br>Kéo thả hoặc bấm · Chọn được nhiều file</div>
        </div>
        <div id="replace-result" style="margin-top:10px"></div>
      </div>
    </div>
  </div>

  <!-- PAGE 3 -->
  <div class="page" id="pg-3">
    <div class="card">
      <div class="card-title">📧 Chọn ticket &amp; gửi Freshdesk <span class="badge badge-mock" style="font-size:11px">Mock API</span></div>

      <!-- Thanh hành động sticky trên đầu -->
      <div id="s3act" style="display:flex;align-items:center;justify-content:space-between;
           flex-wrap:wrap;gap:8px;padding:10px 14px;background:#f8fafc;
           border:1px solid #e5e7eb;border-radius:10px;margin-bottom:14px"></div>

      <div style="display:flex;align-items:flex-end;justify-content:space-between;border-bottom:2px solid #e5e7eb;margin-bottom:16px">
        <div class="tab-bar" style="border-bottom:none;margin-bottom:0">
          <div class="tab active" onclick="switchTab('preview',event)">📧 Preview reply</div>
          <div class="tab" onclick="switchTab('tickets',event)">🎫 Danh sách ticket mock</div>
        </div>
        <button class="btn-ghost" onclick="goTo(2)" style="padding-bottom:12px;font-size:13px">← Quay lại</button>
      </div>
      <div class="tab-content active" id="tab-preview">
        <div id="s3body"></div>
      </div>
      <div class="tab-content" id="tab-tickets">
        <div id="ticketListBody"></div>
      </div>
    </div>
  </div>

  <!-- PAGE 4 -->
  <div class="page" id="pg-4">
    <div class="card">
      <div class="card-title">✅ Kết quả gửi Freshdesk</div>
      <div id="s4act" style="display:flex;gap:8px;flex-wrap:wrap;margin-bottom:16px"></div>
      <div id="s4body"></div>
    </div>
  </div>

</div><!-- /wrap -->

<script>
let selFile = null, res = null, currentPage = 1;

/* ── Wizard navigation ── */
function goTo(n) {
  document.querySelectorAll('.page').forEach(p => p.classList.remove('active'));
  document.getElementById('pg-' + n).classList.add('active');
  for (let i = 1; i <= 4; i++) {
    const si = document.getElementById('si-' + i);
    si.classList.remove('active','done');
    if (i < n)  { si.classList.add('done');   document.getElementById('sn-'+i).textContent = '✓'; }
    if (i === n) { si.classList.add('active'); document.getElementById('sn-'+i).textContent = String(i); }
    if (i > n)  { document.getElementById('sn-'+i).textContent = String(i); }
  }
  // Click trên stepper để quay lại bước đã xong
  for (let i = 1; i < n; i++) {
    const si = document.getElementById('si-' + i);
    si.onclick = () => goTo(i);
  }
  currentPage = n;
  window.scrollTo({top:0, behavior:'smooth'});
}

/* ── Tabs ── */
function switchTab(name, ev) {
  document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
  document.querySelectorAll('.tab-content').forEach(t => t.classList.remove('active'));
  ev.target.classList.add('active');
  document.getElementById('tab-' + name).classList.add('active');
  if (name === 'tickets') loadTickets();
}

/* ── Step 1 ── */
function onFileSelect(input) {
  if (!input.files.length) return;
  selFile = input.files[0];
  document.getElementById('fname').textContent = '📎 ' + selFile.name;
  document.getElementById('btn1').disabled = false;
}

async function processExcel() {
  if (!selFile) return;
  const btn = document.getElementById('btn1');
  btn.disabled = true; btn.textContent = '⏳ Đang xử lý...';
  const prog = document.getElementById('p1');
  const fill = document.getElementById('f1');
  const pt   = document.getElementById('pt1');
  prog.style.display = 'block';
  let pct = 0;
  const iv = setInterval(() => {
    pct = Math.min(pct + 8, 88);
    fill.style.width = pct + '%';
    pt.textContent = pct < 30 ? 'Đọc file Excel...' : pct < 60 ? 'Tạo file Word...' : 'Hoàn thiện...';
  }, 280);

  const fd = new FormData();
  fd.append('file', selFile);
  try {
    const r = await fetch('/process', {method:'POST', body:fd});
    const d = await r.json();
    clearInterval(iv); fill.style.width = '100%'; pt.textContent = 'Xong!';
    if (d.success) { res = d; buildStep2(d); goTo(2); }
    else { pt.textContent = '❌ ' + d.error; btn.disabled = false; btn.textContent = '⚙️ Tạo file Word chứng từ'; }
  } catch(e) {
    clearInterval(iv); pt.textContent = '❌ ' + e.message;
    btn.disabled = false; btn.textContent = '⚙️ Tạo file Word chứng từ';
  }
}

/* ── Step 2 ── */
function buildStep2(d) {
  const ready = d.files.filter(f => f.ready.toUpperCase() === 'YES').length;
  let html = `<div class="stat-row">
    <div class="stat"><div class="num">${d.total_rows}</div><div class="lbl">Giao dịch</div></div>
    <div class="stat"><div class="num">${d.files.length}</div><div class="lbl">File Word tạo ra</div></div>
    <div class="stat"><div class="num">${ready}</div><div class="lbl">Sẵn sàng gửi FD</div></div>
  </div>`;
  html += '<div class="alert alert-success">✅ Tạo thành công ' + d.files.length + ' file Word!</div>';
  html += '<table><thead><tr><th>File Word</th><th>Bank</th><th>Ticket ID</th><th>Số GD</th><th>Trạng thái</th></tr></thead><tbody>';
  for (const f of d.files) {
    const b = f.ready.toUpperCase() === 'YES'
      ? '<span class="badge badge-ok">Sẵn sàng</span>'
      : '<span class="badge badge-skip">Chưa YES</span>';
    html += `<tr><td>📄 ${f.file_name}</td><td>${f.bank}</td><td>#${f.ticket_id||'-'}</td><td>${f.num_rows}</td><td>${b}</td></tr>`;
  }
  html += '</tbody></table>';
  document.getElementById('s2body').innerHTML = html;
  document.getElementById('s2act').innerHTML = `
    <button class="btn btn-gray" onclick="downloadAll()">⬇️ Tải ZIP file Word</button>
    <button class="btn btn-blue" id="btn-next" onclick="convertAndNext(${ready})">Tiếp tục gửi Freshdesk (${ready} ticket) →</button>
  `;
  document.getElementById('convert-progress').style.display = 'none';
  document.getElementById('replace-zone').style.display = 'block';
}

/* ── Convert PDF rồi qua Step 3 ── */
async function convertAndNext(readyCount) {
  const btn = document.getElementById('btn-next');
  btn.disabled = true; btn.textContent = '⏳ Đang chuyển PDF...';
  const prog = document.getElementById('convert-progress');
  const fill = document.getElementById('convert-fill');
  const txt  = document.getElementById('convert-text');
  prog.style.display = 'block';
  // Giả lập progress bar trong khi chờ server
  let pct = 0;
  const iv = setInterval(() => {
    pct = Math.min(pct + 3, 90);
    fill.style.width = pct + '%';
    txt.textContent = `⏳ Đang chuyển đổi sang PDF... (${Math.round(pct)}%)`;
  }, 200);
  try {
    const r = await fetch('/convert_pdf', {method:'POST'});
    const d = await r.json();
    clearInterval(iv);
    if (!d.success) {
      // Không chặn flow — cảnh báo rồi vẫn tiếp tục với file docx gốc
      fill.style.width = '100%'; fill.style.background = '#f59e0b';
      txt.textContent = `⚠️ Không thể chuyển PDF (${d.error}) — sẽ gửi file Word gốc.`;
      await new Promise(r => setTimeout(r, 1800));
      buildStep3(); goTo(3);
      return;
    }
    fill.style.width = '100%'; fill.style.background = '#15803d';
    const failNote = d.failed.length > 0 ? ` (${d.failed.length} file giữ Word)` : '';
    txt.textContent = `✅ Đã chuyển ${d.converted.length} file sang PDF${failNote}`;
    await new Promise(r => setTimeout(r, 700));
    buildStep3(); goTo(3);
  } catch(e) {
    clearInterval(iv);
    fill.style.background = '#ef4444'; fill.style.width = '100%';
    txt.textContent = '❌ Lỗi kết nối: ' + e.message;
    btn.disabled = false; btn.textContent = 'Tiếp tục gửi Freshdesk →';
  }
}

/* ── Step 3 ── */
function buildStep3() {
  const ready = res.files.filter(f => f.ready.toUpperCase() === 'YES');
  let html = '<div class="alert alert-warn" style="margin-bottom:12px">🧪 <b>Mock Mode</b> — Ticket giả lập, không gửi email thật.</div>';
  for (const f of ready) {
    html += `<div class="preview-email ticket-row" id="row-${f.ticket_id}">
      <input type="checkbox" class="ticket-checkbox" data-tid="${f.ticket_id}" checked onchange="onTicketCheck()">
      <div class="meta">
        <b>Ticket #${f.ticket_id}</b> · ${f.file_name}<br>
        To: mock_bank_${String(f.ticket_id).slice(-4)}@mockbank.com · CC: banksupport@vng.com.vn, chargebackzp@vng.com.vn
      </div>
      Dear anh chị,<br><br>
      Zalopay cung cấp chứng từ khiếu nại như đính kèm. Xin cảm ơn.<br><br>
      <span style="color:#003087;font-size:12px"><b>Bui Thi Quynh Hoa (Ms.)</b> · Operations Team · E: Hoabtq@vng.com.vn · P: (028) 3962 3888 ext:3227</span>
      <div><span class="preview-attach">📎 ${f.file_name}</span></div>
    </div>`;
  }
  document.getElementById('s3body').innerHTML = html;
  updateSendBtn();
}

function getCheckedTids() {
  return [...document.querySelectorAll('.ticket-checkbox:checked')].map(c => c.dataset.tid);
}
function getAllTids() {
  return [...document.querySelectorAll('.ticket-checkbox')].map(c => c.dataset.tid);
}
function toggleAll(master) {
  document.querySelectorAll('.ticket-checkbox').forEach(c => {
    c.checked = master.checked;
    document.getElementById('row-' + c.dataset.tid).classList.toggle('unchecked', !master.checked);
  });
  updateSendBtn();
}
function onTicketCheck() {
  document.querySelectorAll('.ticket-checkbox').forEach(c => {
    document.getElementById('row-' + c.dataset.tid).classList.toggle('unchecked', !c.checked);
  });
  updateSendBtn();
}
function updateSendBtn() {
  const checked = getCheckedTids(), all = getAllTids();
  const allChecked = checked.length === all.length;
  const selBtn = allChecked
    ? `<button class="btn btn-gray" onclick="toggleAllBtn(false)" style="padding:7px 14px">☑ Bỏ chọn tất cả</button>`
    : `<button class="btn btn-gray" onclick="toggleAllBtn(true)"  style="padding:7px 14px">☐ Chọn tất cả</button>`;
  const count = `<span style="font-size:13px;font-weight:600;color:#003087">${checked.length}/${all.length} ticket</span>`;
  const send = checked.length > 0
    ? `<button class="btn btn-green" onclick="sendFreshdesk()">🚀 Gửi ${checked.length} reply</button>`
    : `<button class="btn btn-gray" disabled style="cursor:not-allowed">— Chưa chọn ticket nào —</button>`;
  document.getElementById('s3act').innerHTML =
    `<span>${count}</span><span style="display:flex;gap:8px;margin-left:auto">${selBtn}${send}</span>`;
}
function toggleAllBtn(selectAll) {
  document.querySelectorAll('.ticket-checkbox').forEach(c => {
    c.checked = selectAll;
    document.getElementById('row-' + c.dataset.tid).classList.toggle('unchecked', !selectAll);
  });
  updateSendBtn();
}

async function loadTickets() {
  const r = await fetch('/mock-freshdesk/tickets');
  const d = await r.json();
  let html = '<table><thead><tr><th>Ticket ID</th><th>Subject</th><th>CC</th><th>Replies</th></tr></thead><tbody>';
  for (const t of d.tickets) {
    html += `<tr><td>#${t.id}</td><td>${t.subject}</td><td style="font-size:12px">${t.cc_emails.join(', ')}</td><td>${t.replies_count > 0 ? '<span class="badge badge-ok">'+t.replies_count+' reply</span>' : '<span class="badge badge-skip">Chưa có</span>'}</td></tr>`;
  }
  html += '</tbody></table>';
  document.getElementById('ticketListBody').innerHTML = html;
}

/* ── Step 4 ── */
async function sendFreshdesk() {
  const btn = event.target;
  btn.disabled = true; btn.textContent = '⏳ Đang gửi...';
  const r = await fetch('/send_freshdesk', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({ticket_ids: getCheckedTids()})
  });
  const d = await r.json();
  const ok = d.results.filter(x => x.status.includes('SUCCESS')).length;
  const sk = d.results.filter(x => x.status.startsWith('SKIPPED')).length;
  let html = `<div class="stat-row">
    <div class="stat"><div class="num" style="color:#15803d">${ok}</div><div class="lbl">Gửi thành công</div></div>
    <div class="stat"><div class="num" style="color:#92400e">${sk}</div><div class="lbl">Bỏ qua</div></div>
    <div class="stat"><div class="num">${d.results.length}</div><div class="lbl">Tổng</div></div>
  </div>`;
  html += '<table><thead><tr><th>Ticket ID</th><th>File đính kèm</th><th>CC sau khi lọc</th><th>Kết quả</th></tr></thead><tbody>';
  for (const x of d.results) {
    const b = x.status.includes('SUCCESS')
      ? `<span class="badge badge-ok">✓ ${x.status}</span>`
      : `<span class="badge badge-skip">${x.status}</span>`;
    const fileCell = x.status.includes('SUCCESS')
      ? `<a href="/view/${encodeURIComponent(x.file_name)}" target="_blank"
            style="color:#0369a1;text-decoration:none;font-size:12px" title="Xem chứng từ">📎 ${x.file_name}</a>`
      : `<span style="font-size:12px;color:#9ca3af">📎 ${x.file_name}</span>`;
    html += `<tr><td>#${x.ticket_id}</td><td>${fileCell}</td><td style="font-size:11px">${x.cc.join('<br>')}</td><td>${b}</td></tr>`;
  }
  html += '</tbody></table>';
  document.getElementById('s4act').innerHTML = `
    <button class="btn btn-gray" onclick="downloadLog()">⬇️ Tải log kết quả (.xlsx)</button>
    <button class="btn-ghost" onclick="resetAll()">↩ Bắt đầu lại</button>
  `;
  document.getElementById('s4body').innerHTML = html;
  goTo(4);
  await loadTickets();
}

/* ── Replace upload ── */
function handleReplaceDrop(event) {
  event.preventDefault();
  document.getElementById('replace-drop').style.borderColor = '#cbd5e1';
  uploadReplaceFiles(event.dataTransfer.files);
}
function handleReplaceSelect(input) {
  if (input.files.length) uploadReplaceFiles(input.files);
}
async function uploadReplaceFiles(files) {
  const fd = new FormData();
  for (const f of files) fd.append('files[]', f);
  const div = document.getElementById('replace-result');
  div.innerHTML = '<span style="color:#64748b;font-size:13px">⏳ Đang upload...</span>';
  try {
    const r = await fetch('/upload_replace', {method:'POST', body:fd});
    const d = await r.json();
    let html = '';
    if (d.replaced.length)
      html += '<div class="alert alert-success">✅ Đã thay thế ' + d.replaced.length + ' file: ' +
              d.replaced.map(n=>`<span class="preview-attach" style="margin:2px 3px">📄 ${n}</span>`).join('') + '</div>';
    if (d.skipped.length)
      html += '<div class="alert alert-warn">⚠️ Bỏ qua: ' + d.skipped.map(s=>s.name+' ('+s.reason+')').join(', ') + '</div>';
    div.innerHTML = html || '<span style="color:#64748b;font-size:13px">Không có file nào được xử lý.</span>';
  } catch(e) {
    div.innerHTML = '<div class="alert" style="background:#fef2f2;color:#7f1d1d">❌ ' + e.message + '</div>';
  }
}

/* ── Utils ── */
function downloadAll() { window.location.href = '/download_zip'; }
function downloadLog() { window.location.href = '/download_log'; }
function resetAll() {
  selFile = null; res = null;
  document.getElementById('fname').textContent = '';
  document.getElementById('btn1').disabled = true;
  document.getElementById('btn1').textContent = '⚙️ Tạo file Word chứng từ';
  document.getElementById('p1').style.display = 'none';
  document.getElementById('f1').style.width = '0%';
  document.getElementById('replace-zone').style.display = 'none';
  document.getElementById('replace-result').innerHTML = '';
  goTo(1);
}
</script>
</body>
</html>
"""

# =====================
# ROUTES
# =====================
@app.route("/")
def index():
    return render_template_string(HTML)

@app.route("/process", methods=["POST"])
def process():
    try:
        if "file" not in request.files:
            return jsonify({"success": False, "error": "Không có file upload"})

        file = request.files["file"]
        excel_path = UPLOAD_FOLDER / "input.xlsx"
        file.save(excel_path)

        # Clear old outputs
        shutil.rmtree(OUTPUT_FOLDER, ignore_errors=True)
        OUTPUT_FOLDER.mkdir(exist_ok=True)

        df = pd.read_excel(excel_path).dropna(how="all")
        if GROUP_BY_COLUMN not in df.columns:
            return jsonify({"success": False, "error": f"Không có cột '{GROUP_BY_COLUMN}' trong file Excel"})

        # Load ticket mapping tu reply list
        ticket_map = {}
        if FRESHDESK_AUTO.exists():
            df_fd = pd.read_excel(FRESHDESK_AUTO)
            for _, row in df_fd.iterrows():
                fname = str(row.get("file_name", "")).strip()
                tid   = str(row.get("ticket_id", "")).strip()
                ready = str(row.get("ready_to_send", "")).strip()
                if fname:
                    ticket_map[fname] = {"ticket_id": tid, "ready": ready}

        result_files = []
        freshdesk_data = []
        used_names = defaultdict(int)

        for group_val, group_df in df.groupby(GROUP_BY_COLUMN, dropna=False, sort=False):
            first = group_df.iloc[0]
            ctx   = build_context(first)
            ctx["document_title"] = build_doc_title(ctx)
            ctx["service_desc_1"] = get_service_line1(group_df)
            ctx["service_desc_2"] = get_service_line2(group_df)
            ctx["wallet_balance_note"] = get_wallet_note(first)

            ten_file = ctx.get("ten_file", "")
            fd_info  = ticket_map.get(ten_file, {})
            ticket_id    = fd_info.get("ticket_id", "")
            ready_to_send = fd_info.get("ready", "")

            raw_name  = build_output_filename(ctx)
            file_name = safe_filename(raw_name)
            stem, ext = os.path.splitext(file_name)
            used_names[file_name] += 1
            if used_names[file_name] > 1:
                file_name = f"{stem}_{used_names[file_name]}{ext}"

            out_path = OUTPUT_FOLDER / file_name

            # Tao Word
            if WORD_TEMPLATE.exists():
                from docx import Document
                remove_mail_merge(WORD_TEMPLATE)
                doc = Document(WORD_TEMPLATE)
                replace_doc(doc, ctx)
                render_rows(doc, group_df)
                doc.save(out_path)
                remove_mail_merge(out_path)
            else:
                # Fallback: tao Word don gian
                from docx import Document
                doc = Document()
                doc.add_heading(ctx.get("document_title", file_name), 0)
                doc.add_paragraph(f"Bank: {ctx.get('Bank','')}")
                doc.add_paragraph(f"Mô tả: {ctx.get('service_desc_1','')}")
                t = doc.add_table(rows=1, cols=4)
                t.style = "Table Grid"
                for i, h in enumerate(["BC Trace No","Số tiền","Thời gian","Trạng thái"]):
                    t.rows[0].cells[i].text = h
                for _, row in group_df.iterrows():
                    c = build_context(row)
                    r = t.add_row().cells
                    r[0].text = c.get("bc_trace_no","")
                    r[1].text = c.get("amount","")
                    r[2].text = c.get("trans_time","")
                    r[3].text = c.get("bc_status","")
                doc.save(out_path)

            result_files.append({
                "file_name": file_name,
                "bank": ctx.get("Bank",""),
                "user_id": clean_value(group_val),
                "num_rows": len(group_df),
                "ticket_id": ticket_id,
                "ready": ready_to_send,
                "path": str(out_path),
            })
            freshdesk_data.append({
                "ticket_id": ticket_id,
                "ready_to_send": ready_to_send,
                "file_name": file_name,
                "bank": ctx.get("Bank",""),
                "attachment_path": str(out_path.resolve()),
            })

        state["files"]          = result_files
        state["freshdesk_data"] = freshdesk_data

        return jsonify({
            "success": True,
            "files": result_files,
            "total_rows": len(df),
        })
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e) + "\n" + traceback.format_exc()})


@app.route("/send_freshdesk", methods=["POST"])
def send_freshdesk():
    import requests as req_lib
    results = []
    base_url = "http://localhost:8080/mock-freshdesk"

    # Lấy danh sách ticket_id được chọn từ frontend (nếu có)
    body = request.get_json(silent=True) or {}
    selected_ids = set(str(x).strip() for x in body.get("ticket_ids", []))

    for item in state.get("freshdesk_data", []):
        tid   = str(item.get("ticket_id","")).strip()
        ready = str(item.get("ready_to_send","")).strip().upper()
        fname = item.get("file_name","")
        apath = Path(item.get("attachment_path",""))

        # Bỏ qua nếu: chưa ready, hoặc không nằm trong danh sách được chọn
        if ready != "YES":
            results.append({"ticket_id": tid, "file_name": fname, "cc": [], "status": "SKIPPED (not ready)"})
            continue
        if selected_ids and tid not in selected_ids:
            results.append({"ticket_id": tid, "file_name": fname, "cc": [], "status": "SKIPPED (bỏ chọn)"})
            continue

        # Lay CC tu mock API
        try:
            r = req_lib.get(f"{base_url}/api/v2/tickets/{tid}", timeout=5)
            cc_raw = r.json().get("cc_emails", []) if r.status_code == 200 else []
        except Exception:
            cc_raw = []

        # Loc va them CC
        cc_cleaned = []
        for email in cc_raw:
            email = str(email).strip().lower()
            if "@" in email and email.split("@")[-1] not in REMOVE_DOMAINS:
                cc_cleaned.append(email)
        cc_final = list(dict.fromkeys(cc_cleaned + ADD_CC_EMAILS))

        # Gui reply qua mock API
        try:
            data = [("body", """<div>Dear anh chị,<br><br>Zalopay cung cấp chứng từ khiếu nại như đính kèm. Xin cảm ơn.<br><br>Bui Thi Quynh Hoa (Ms.) · Operations Team</div>""")]
            for email in cc_final:
                data.append(("cc_emails[]", email))

            if apath.exists():
                with open(apath, "rb") as f:
                    files = {"attachments[]": (apath.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                    resp = req_lib.post(f"{base_url}/api/v2/tickets/{tid}/reply", data=data, files=files, timeout=10)
            else:
                resp = req_lib.post(f"{base_url}/api/v2/tickets/{tid}/reply", data=data, timeout=10)

            status = "SUCCESS (MOCK)" if resp.status_code == 201 else f"FAILED ({resp.status_code})"
        except Exception as e:
            status = f"ERROR: {str(e)}"

        results.append({"ticket_id": tid, "file_name": fname, "cc": cc_final, "status": status})
        time.sleep(0.2)

    state["send_results"] = results
    pd.DataFrame(results).to_excel(LOG_FILE, index=False)
    return jsonify({"results": results})


@app.route("/upload_replace", methods=["POST"])
def upload_replace():
    """Nhận file .docx upload, ghi đè vào output_grouped theo tên file."""
    uploaded_files = request.files.getlist("files[]")
    replaced = []
    skipped = []
    for f in uploaded_files:
        fname = Path(f.filename).name  # Lấy tên file, bỏ đường dẫn
        if not fname.lower().endswith(".docx"):
            skipped.append({"name": fname, "reason": "Không phải file .docx"})
            continue
        target = OUTPUT_FOLDER / fname
        f.save(target)
        # Cập nhật attachment_path trong state nếu file đã tồn tại
        for item in state.get("freshdesk_data", []):
            if item.get("file_name") == fname:
                item["attachment_path"] = str(target.resolve())
        replaced.append(fname)
    return jsonify({"replaced": replaced, "skipped": skipped})


def _do_convert_batch():
    """
    Batch-convert tất cả .docx trong OUTPUT_FOLDER → PDF.
    - Windows : dùng docx2pdf (Microsoft Word COM) — 1 lần mở Word cho cả batch
    - Linux/Mac: dùng LibreOffice headless              — 1 lần khởi động cho cả batch
    Trả về (ok: bool, error_msg: str)
    """
    import platform
    docx_files = list(OUTPUT_FOLDER.glob("*.docx"))
    if not docx_files:
        return True, ""

    if platform.system() == "Windows":
        try:
            from docx2pdf import convert as docx2pdf_convert
        except ImportError:
            return False, "Thiếu docx2pdf. Chạy: python -m pip install docx2pdf"
        try:
            docx2pdf_convert(str(OUTPUT_FOLDER), str(OUTPUT_FOLDER))
            return True, ""
        except Exception as e:
            return False, str(e)
    else:
        # Linux / Mac → LibreOffice headless
        import subprocess, shutil
        lo = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo:
            return False, "LibreOffice chưa cài. Chạy: sudo apt-get install -y libreoffice"
        try:
            cmd = [lo, "--headless", "--convert-to", "pdf",
                   "--outdir", str(OUTPUT_FOLDER)] + [str(f) for f in docx_files]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
            if result.returncode != 0:
                return False, result.stderr[:300]
            return True, ""
        except subprocess.TimeoutExpired:
            return False, "Timeout: convert mất quá lâu (>180s)"
        except Exception as e:
            return False, str(e)


@app.route("/convert_pdf", methods=["POST"])
def convert_pdf():
    """Batch-convert toàn bộ OUTPUT_FOLDER → PDF, tự nhận diện Windows/Linux."""
    ok, err = _do_convert_batch()
    if not ok:
        return jsonify({"success": False, "error": err}), 500

    # Cập nhật state sau khi convert xong
    converted, failed = [], []
    for item in state.get("freshdesk_data", []):
        if str(item.get("ready_to_send", "")).strip().upper() != "YES":
            continue
        docx_path = Path(item.get("attachment_path", ""))
        pdf_path  = docx_path.with_suffix(".pdf")
        if pdf_path.exists():
            item["attachment_path"] = str(pdf_path)
            item["file_name"]       = pdf_path.name
            converted.append(pdf_path.name)
        else:
            failed.append({"name": docx_path.name, "reason": "PDF không được tạo"})

    for f in state.get("files", []):
        pdf_name = Path(f["file_name"]).with_suffix(".pdf").name
        if (OUTPUT_FOLDER / pdf_name).exists():
            f["file_name"] = pdf_name

    return jsonify({"success": True, "converted": converted, "failed": failed})


@app.route("/download_zip")
def download_zip():
    zip_path = BASE_DIR / "output.zip"
    with zipfile.ZipFile(zip_path, "w") as zf:
        for f in OUTPUT_FOLDER.glob("*.docx"):
            zf.write(f, f.name)
    return send_file(zip_path, as_attachment=True, download_name="chung_tu_zalopay.zip")


@app.route("/view/<path:filename>")
def view_docx(filename):
    """Xem chứng từ: ưu tiên PDF (Chrome native), fallback mammoth HTML."""
    # Nếu là PDF → trả về thẳng cho Chrome render
    filepath = OUTPUT_FOLDER / filename
    if filename.lower().endswith(".pdf") and filepath.exists():
        return send_file(filepath, mimetype="application/pdf")

    # Nếu là docx nhưng đã có bản PDF cùng tên → dùng PDF
    if filename.lower().endswith(".docx"):
        pdf_path = OUTPUT_FOLDER / Path(filename).with_suffix(".pdf").name
        if pdf_path.exists():
            return send_file(pdf_path, mimetype="application/pdf")

    # Fallback: mammoth → HTML
    if not filepath.exists():
        return "<h3>File không tồn tại hoặc đã bị xóa.</h3>", 404
    try:
        import mammoth
    except ImportError:
        return (
            "<h3>Thiếu thư viện <code>mammoth</code>.</h3>"
            "<p>Chạy: <code>py -m pip install mammoth</code></p>"
        ), 500
    try:
        with open(filepath, "rb") as f:
            result = mammoth.convert_to_html(f)
        body = result.value
        page = f"""<!DOCTYPE html>
<html lang="vi"><head><meta charset="UTF-8"><title>{filename}</title>
<style>
  *{{box-sizing:border-box}}
  body{{font-family:'Times New Roman',serif;font-size:12pt;line-height:1.6;
       margin:0;padding:0;color:#111;background:#f5f5f5}}
  .docbar{{background:#003087;color:white;padding:11px 28px;
           display:flex;align-items:center;justify-content:space-between;
           font-size:12pt;position:sticky;top:0;z-index:10}}
  .docbar a{{color:#93c5fd;text-decoration:none;font-size:11pt}}
  .doccontent{{background:white;padding:32px 40px;margin:20px auto;
              width:fit-content;min-width:600px;max-width:98vw;
              box-shadow:0 2px 10px rgba(0,0,0,.1);min-height:80vh;overflow-x:auto}}
  table{{border-collapse:collapse;margin:8px 0;white-space:nowrap}}
  td,th{{border:1px solid #888;padding:5px 9px;font-size:10.5pt}}
  th{{background:#f0f0f0;font-weight:600;text-align:center}}
  td{{vertical-align:top}} p{{margin:4px 0}}
  @media print{{.docbar{{display:none}}.doccontent{{box-shadow:none;padding:0;margin:0;width:100%}}table{{white-space:normal}}}}
</style></head><body>
<div class="docbar">
  <span>📄 {filename}</span>
  <span style="display:flex;gap:20px">
    <a href="/download_single/{filename}">⬇️ Tải về</a>
    <a href="javascript:window.print()">🖨️ In</a>
  </span>
</div>
<div class="doccontent">{body}</div>
</body></html>"""
        return page, 200, {"Content-Type": "text/html; charset=utf-8"}
    except Exception as e:
        import traceback
        return f"<pre>Lỗi: {e}\n{traceback.format_exc()}</pre>", 500


@app.route("/download_single/<path:filename>")
def download_single(filename):
    filepath = OUTPUT_FOLDER / filename
    if not filepath.exists():
        return "File không tồn tại", 404
    return send_file(filepath, as_attachment=True, download_name=filename)


@app.route("/download_log")
def download_log():
    if LOG_FILE.exists():
        return send_file(LOG_FILE, as_attachment=True, download_name="freshdesk_log.xlsx")
    return "Chưa có log", 404


@app.route("/health")
def health():
    return jsonify({"status": "ok", "mock_tickets": len(MOCK_DB), "mode": "MOCK"})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    print(f"🚀 ZOP_HR_INBOX chạy tại http://localhost:{port}")
    print(f"🎫 Mock Freshdesk tickets: {len(MOCK_DB)}")
    app.run(host="0.0.0.0", port=port, debug=False)
